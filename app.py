"""
Research Paper Writing Agent
A persistent, style-aware academic writing assistant built with Streamlit + Anthropic Claude.
Run: streamlit run app.py
Requires: ANTHROPIC_API_KEY environment variable
"""

import io
import json
import re
import uuid
from datetime import datetime
from pathlib import Path

import anthropic
import requests
import streamlit as st
from bs4 import BeautifulSoup
from dotenv import load_dotenv

load_dotenv()

try:
    import pypdf
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    from fpdf import FPDF
    FPDF_SUPPORT = True
except ImportError:
    FPDF_SUPPORT = False

try:
    import openpyxl
    XLSX_SUPPORT = True
except ImportError:
    XLSX_SUPPORT = False

# ==============================================================================
# PAPER_CONTEXT — global defaults, overridable per-session in the UI
# ==============================================================================
PAPER_CONTEXT = {
    "paper_title": "",           # e.g., "Attention Is All You Need"
    "research_question": "",     # The main research question or hypothesis
    "models_methods": "",        # Models or methods used
    "dataset_description": "",   # Dataset name, size, and key properties
    "key_findings": "",          # Primary results and contributions
    "target_journal": "",        # Target venue (e.g., "NeurIPS 2025")
    "co_authors": "",            # Author names, separated by commas
}
# ==============================================================================

SESSIONS_FILE = "sessions.json"
RULES_FILE = "agent_rules.json"
COUNCIL_RULES_FILE = "council_rules.json"
COUNCIL_MEMORY_FILE = "council_memory.txt"
RESEARCH_FILE = "research_notes.txt"
MODEL = "claude-sonnet-4-6"
HAIKU_MODEL = "claude-haiku-4-5-20251001"
SECTIONS = ["Introduction", "Methods", "Results", "Discussion", "Conclusion", "Abstract", "Title"]
# Max characters of document content stored per document
DOC_CONTENT_LIMIT = 25_000

_CONTEXT_FIELDS = [
    ("paper_title",         "Paper Title",       "e.g., Attention Is All You Need"),
    ("research_question",   "Research Question", "The main research question or hypothesis"),
    ("models_methods",      "Models / Methods",  "Models or methods used"),
    ("dataset_description", "Dataset",           "Dataset name, size, and key properties"),
    ("key_findings",        "Key Findings",      "Primary results and contributions"),
    ("target_journal",      "Target Venue",      "e.g., NeurIPS 2025"),
    ("co_authors",          "Co-authors",        "Author names, separated by commas"),
]


# ── Document fetching ──────────────────────────────────────────────────────────

def extract_pdf_text(pdf_bytes: bytes) -> str:
    if not PDF_SUPPORT:
        raise RuntimeError("pypdf not installed")
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip())


def extract_file_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch file extraction based on extension. Returns plain text."""
    import csv as _csv
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return extract_pdf_text(file_bytes)

    if ext in (".txt", ".md", ".tex", ".rst"):
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return file_bytes.decode(enc)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="replace")

    if ext == ".csv":
        text = file_bytes.decode("utf-8", errors="replace")
        reader = _csv.reader(text.splitlines())
        rows = list(reader)
        if not rows:
            return ""
        col_widths = [max(len(str(r[i])) for r in rows if i < len(r)) for i in range(len(rows[0]))]
        lines = []
        for row in rows:
            parts = [str(cell).ljust(col_widths[i]) for i, cell in enumerate(row) if i < len(col_widths)]
            lines.append("  ".join(parts))
        return "\n".join(lines)

    if ext in (".xlsx", ".xls"):
        if not XLSX_SUPPORT:
            raise RuntimeError("openpyxl not installed — run: pip install openpyxl")
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                if any(c is not None for c in row):
                    parts.append("\t".join("" if c is None else str(c) for c in row))
        return "\n".join(parts)

    if ext == ".docx":
        if not DOCX_SUPPORT:
            raise RuntimeError("python-docx not installed — run: pip install python-docx")
        doc = DocxDocument(io.BytesIO(file_bytes))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    raise ValueError(f"Unsupported file type: {ext}")


def read_local_path(path: str) -> tuple[str, str]:
    """Read a file from the local filesystem by path. Returns (name, text)."""
    p = Path(path.strip())
    if not p.exists():
        raise FileNotFoundError(f"File not found: {p}")
    if not p.is_file():
        raise ValueError(f"Path is not a file: {p}")
    content = extract_file_text(p.read_bytes(), p.name)
    return p.name, content


def _docx_add_inline(para, text: str) -> None:
    """Parse **bold** and *italic* markdown into docx runs."""
    pattern = re.compile(r'(\*\*[^*\n]+?\*\*|\*[^*\n]+?\*|[^*]+)')
    for match in pattern.finditer(text):
        chunk = match.group(0)
        if chunk.startswith('**') and chunk.endswith('**'):
            run = para.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith('*') and chunk.endswith('*'):
            run = para.add_run(chunk[1:-1])
            run.italic = True
        else:
            para.add_run(chunk)


@st.cache_data(show_spinner=False)
def generate_docx(draft_content: str, paper_title: str, section: str) -> bytes:
    """Generate a Word document from a markdown-formatted draft."""
    if not DOCX_SUPPORT:
        raise RuntimeError("python-docx not installed — run: pip install python-docx")
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.25)
        sec.right_margin = Inches(1.25)

    # Title block
    title_para = doc.add_heading(level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run(paper_title)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f"{section} Section")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    doc.add_paragraph()

    for line in draft_content.split('\n'):
        stripped = line.rstrip()
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif not stripped:
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            _docx_add_inline(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@st.cache_data(show_spinner=False)
def generate_pdf(draft_content: str, paper_title: str, section: str) -> bytes:
    """Generate a PDF from a markdown-formatted draft using fpdf2."""
    if not FPDF_SUPPORT:
        raise RuntimeError("fpdf2 not installed — run: pip install fpdf2")

    pdf = FPDF()
    pdf.set_margins(25, 25, 25)
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()

    # Title block
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, paper_title[:100], align='C')
    pdf.set_font("Helvetica", "I", 12)
    pdf.multi_cell(0, 8, f"{section} Section", align='C')
    pdf.ln(8)

    for line in draft_content.split('\n'):
        stripped = line.rstrip()
        # Strip inline markdown for PDF plain-text rendering
        clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', stripped)
        clean = re.sub(r'\*([^*]+)\*', r'\1', clean)

        if stripped.startswith('### '):
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 6, clean[4:])
            pdf.set_font("Helvetica", "", 11)
        elif stripped.startswith('## '):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 7, clean[3:])
            pdf.set_font("Helvetica", "", 11)
            pdf.ln(1)
        elif stripped.startswith('# '):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, clean[2:])
            pdf.set_font("Helvetica", "", 11)
            pdf.ln(2)
        elif not stripped:
            pdf.ln(4)
        else:
            pdf.set_font("Helvetica", "", 11)
            try:
                pdf.multi_cell(0, 6, clean)
            except Exception:
                safe = clean.encode('latin-1', errors='replace').decode('latin-1')
                pdf.multi_cell(0, 6, safe)

    return bytes(pdf.output())


def fetch_url_content(url: str) -> tuple[str, str]:
    """
    Fetch content from a URL. Returns (document_name, text_content).
    Handles Google Docs, Google Drive files, direct PDF URLs, and regular web pages.
    """
    url = url.strip()

    # Google Docs shared link → export as plain text
    gdoc_match = re.search(r"docs\.google\.com/document/d/([^/?\s]+)", url)
    if gdoc_match:
        doc_id = gdoc_match.group(1)
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
        resp = requests.get(export_url, timeout=30, allow_redirects=True)
        resp.raise_for_status()
        ctype = resp.headers.get("content-type", "")
        # Google redirects to an HTML login page when the doc is private
        if "text/html" in ctype or "accounts.google.com" in resp.url:
            raise ValueError(
                "Could not access this Google Doc — it may be private. "
                "Open the doc, click Share → change to 'Anyone with the link can view', then try again."
            )
        text = resp.text.strip()
        if not text:
            raise ValueError("Google Doc exported successfully but returned empty content.")
        return f"Google Doc ({doc_id[:12]})", text

    # Google Drive file link → try to download
    gdrive_match = re.search(r"drive\.google\.com/file/d/([^/?\s]+)", url)
    if gdrive_match:
        file_id = gdrive_match.group(1)
        download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
        session_req = requests.Session()
        resp = session_req.get(download_url, timeout=30, allow_redirects=True)
        resp.raise_for_status()
        ctype = resp.headers.get("content-type", "")
        # Detect login redirect (private file)
        if "accounts.google.com" in resp.url:
            raise ValueError(
                "Could not access this Google Drive file — it may be private. "
                "Open it in Drive, click Share → change to 'Anyone with the link can view', then try again."
            )
        # Handle virus-scan confirmation page for large files
        if "text/html" in ctype:
            confirm_match = re.search(r'confirm=([0-9A-Za-z_\-]+)', resp.text)
            if confirm_match:
                confirm_token = confirm_match.group(1)
                resp = session_req.get(
                    f"https://drive.google.com/uc?export=download&confirm={confirm_token}&id={file_id}",
                    timeout=60,
                    allow_redirects=True,
                )
                resp.raise_for_status()
                ctype = resp.headers.get("content-type", "")
            else:
                raise ValueError(
                    "Could not download this Google Drive file. "
                    "Make sure it is shared as 'Anyone with the link can view'."
                )
        if "pdf" in ctype:
            return f"Google Drive PDF ({file_id[:12]})", extract_pdf_text(resp.content)
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        return f"Google Drive file ({file_id[:12]})", soup.get_text("\n", strip=True)

    # Fetch the URL
    headers = {"User-Agent": "Mozilla/5.0 (compatible; PaperAgent/1.0)"}
    resp = requests.get(url, timeout=30, headers=headers)
    resp.raise_for_status()
    ctype = resp.headers.get("content-type", "")

    # Direct PDF URL
    if "pdf" in ctype or url.lower().endswith(".pdf"):
        name = url.rstrip("/").split("/")[-1] or "document.pdf"
        return name, extract_pdf_text(resp.content)

    # Regular web page
    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    text = soup.get_text("\n", strip=True)
    text = re.sub(r"\n{3,}", "\n\n", text)
    title_tag = soup.find("title")
    name = title_tag.get_text(strip=True)[:80] if title_tag else url.split("/")[-1] or url
    return name, text


# ── System prompt ──────────────────────────────────────────────────────────────

def build_system_prompt(session: dict | None = None, rules: list | None = None,
                        research_notes: str = "", related_sessions: list | None = None) -> str:
    prompt = ""

    # 1. Permanent learned preferences — highest priority, always first
    if rules:
        prompt += "LEARNED AUTHOR PREFERENCES — ALWAYS APPLY THESE:\n"
        for i, rule in enumerate(rules, 1):
            prompt += f"{i}. {rule['text']}\n"
        prompt += "\nThese reflect what the author has consistently asked for. Never violate them.\n\n---\n\n"

    # 2. Persistent research knowledge — injected before style guide so it's always available
    if research_notes and research_notes.strip():
        prompt += "RESEARCH KNOWLEDGE BASE — know this thoroughly, reference it when writing:\n\n"
        prompt += research_notes.strip()
        prompt += "\n\n---\n\n"

    prompt += """\
You are a specialized agricultural economics researcher and academic writer. Your expertise \
is in beef production economics, feedlot management, enterprise budgeting, and applied \
production economics. You are familiar with the standards of top agricultural economics \
journals: Applied Economic Perspectives and Policy (AEPP), the Journal of Agricultural and \
Resource Economics (JARE), Agribusiness, and the American Journal of Agricultural Economics \
(AJAE).

YOUR JOB: Draft and revise sections of a peer-reviewed research paper. Before writing \
anything, silently reason through: (1) What is the core claim this section must establish? \
(2) What specific evidence, numbers, or comparisons support it? (3) What limitations or \
caveats must be honestly acknowledged? Then write — do not show this reasoning in your output.

DOMAIN KNOWLEDGE — apply this to every section:
- Ag econ papers ground every claim in specific numbers. Vague language like "substantially \
higher costs" is unacceptable; write "feeder purchase costs averaging $X/cwt represented Y% \
of total enterprise cost."
- Break-even analysis requires explicit formula exposition, variable definitions, and \
sensitivity discussion. Readers must be able to replicate calculations.
- Results sections must reference specific table or figure numbers. Never describe a result \
without tying it to evidence.
- The Methods section must justify model specification choices, not just describe them. \
Explain WHY each variable is included.
- Cite prior work in (Author, Year) format. Flag where a citation is needed with [CITE NEEDED].
- Policy implications belong in Discussion, not Results.
- Acknowledge data limitations directly — reviewers at AEPP and JARE will flag them anyway.

SECTION CONVENTIONS:
- Introduction: establish economic significance with numbers, identify the research gap \
precisely, state the contribution in one direct sentence, preview the structure briefly.
- Methods: past tense. Justify every modeling choice. Define every variable with units.
- Results: present tense. Lead with the most important finding. Every claim ties to a table/figure.
- Discussion: interpret, compare to prior literature, state limitations, give policy implications.
- Conclusion: no new information. Restate contribution, summarize key numbers, give actionable takeaway.
- Abstract: one sentence per section function. State the key finding with a number.

WRITING STYLE:
- Direct, confident sentences. No fluff, no over-hedging.
- Mix a high-level claim with a specific grounded number right after it.
- Short punchy sentences after complex ones for emphasis. One idea per sentence.
- Acknowledge tradeoffs honestly — do not gloss over weaknesses.
- Group related items in a natural run, not bullet points.

NEVER USE:
- "It is important to note that..." / "It is worth mentioning..."
- "Furthermore," / "Moreover," / "Additionally," / "Notably," as sentence starters
- Restating the conclusion at the end of every paragraph
- "While X, it is also true that Y" balanced constructions
- Passive voice overuse (active voice preferred throughout)
- Summary sentences that just repeat what was said
- LaTeX math notation ($...$, _{}, ^{}) — write equations as plain text (e.g. BEP = TC / TFW x 100 + FB)\
"""

    # Per-session context takes precedence over global PAPER_CONTEXT
    ctx = (session or {}).get("paper_context") or PAPER_CONTEXT
    if any(v.strip() for v in ctx.values()):
        prompt += "\n\nPAPER CONTEXT:\n"
        if ctx.get("paper_title"):         prompt += f"Title: {ctx['paper_title']}\n"
        if ctx.get("research_question"):   prompt += f"Research question: {ctx['research_question']}\n"
        if ctx.get("models_methods"):      prompt += f"Models/methods: {ctx['models_methods']}\n"
        if ctx.get("dataset_description"): prompt += f"Dataset: {ctx['dataset_description']}\n"
        if ctx.get("key_findings"):        prompt += f"Key findings: {ctx['key_findings']}\n"
        if ctx.get("target_journal"):      prompt += f"Target venue: {ctx['target_journal']}\n"
        if ctx.get("co_authors"):          prompt += f"Co-authors: {ctx['co_authors']}\n"

    # Attached documents
    if session and session.get("documents"):
        prompt += "\n\nREFERENCE DOCUMENTS (use these as source material when writing):\n"
        for i, doc in enumerate(session["documents"], 1):
            content = doc["content"]
            truncated = len(content) > DOC_CONTENT_LIMIT
            if truncated:
                content = content[:DOC_CONTENT_LIMIT]
            prompt += f"\n--- Document {i}: {doc['name']} ---\n{content}\n"
            if truncated:
                prompt += "[...document truncated for length]\n"

    # Cross-session context: other sections of the same paper
    CROSS_SESSION_LIMIT = 3000  # chars per sibling section
    if related_sessions:
        prompt += "\n\nOTHER SECTIONS OF THIS PAPER (already drafted — match their voice, facts, and terminology):\n"
        for s in related_sessions:
            sec = s.get("section", "Unknown")
            draft = s.get("current_draft", "")
            truncated = len(draft) > CROSS_SESSION_LIMIT
            snippet = draft[:CROSS_SESSION_LIMIT]
            prompt += f"\n--- {sec} ---\n{snippet}"
            if truncated:
                prompt += "\n[...truncated]"
            prompt += "\n"

    return prompt


# ── Research notes persistence ─────────────────────────────────────────────────

def load_research_notes() -> str:
    path = Path(RESEARCH_FILE)
    if path.exists():
        try:
            return path.read_text(encoding="utf-8")
        except IOError:
            return ""
    return ""


def save_research_notes(text: str) -> None:
    Path(RESEARCH_FILE).write_text(text, encoding="utf-8")


# ── Rules persistence ──────────────────────────────────────────────────────────

def load_rules() -> list:
    path = Path(RULES_FILE)
    if path.exists():
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return []
    return []


def save_rules(rules: list) -> None:
    with open(RULES_FILE, "w", encoding="utf-8") as f:
        json.dump(rules, f, indent=2, ensure_ascii=False)


# ── Council rules + memory persistence ─────────────────────────────────────────

def load_council_rules() -> list:
    path = Path(COUNCIL_RULES_FILE)
    if path.exists():
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return []
    return []


def save_council_rules(rules: list) -> None:
    with open(COUNCIL_RULES_FILE, "w", encoding="utf-8") as f:
        json.dump(rules, f, indent=2, ensure_ascii=False)


def load_council_memory() -> str:
    path = Path(COUNCIL_MEMORY_FILE)
    if path.exists():
        try:
            return path.read_text(encoding="utf-8")
        except IOError:
            return ""
    return ""


def save_council_memory(text: str) -> None:
    Path(COUNCIL_MEMORY_FILE).write_text(text, encoding="utf-8")


# ── Session persistence ────────────────────────────────────────────────────────

def load_sessions() -> dict:
    path = Path(SESSIONS_FILE)
    if path.exists():
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return {}
    return {}


def save_sessions(sessions: dict) -> None:
    with open(SESSIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(sessions, f, indent=2, ensure_ascii=False)


def new_session(paper_title: str, section: str) -> tuple[str, dict]:
    session_id = str(uuid.uuid4())[:8]
    session = {
        "session_id": session_id,
        "timestamp": datetime.now().isoformat(),
        "last_modified": datetime.now().isoformat(),
        "paper_title": paper_title,
        "section": section,
        "paper_context": {key: "" for key, *_ in _CONTEXT_FIELDS},
        "conversation_history": [],
        "current_draft": "",
        "drafts": [],
        "iteration_notes": [],
        "documents": [],
        "council_reviews": {},
    }
    return session_id, session


def touch_session(session: dict) -> None:
    """Update last_modified timestamp."""
    session["last_modified"] = datetime.now().isoformat()


# ── Claude API ─────────────────────────────────────────────────────────────────

def truncate_history(history: list, max_pairs: int = 4) -> list:
    """Keep conversation history within context limits.

    Preserves the first pair (original notes → V1) for intent continuity,
    then the most recent (max_pairs - 1) exchange pairs. The current draft
    is always in the last assistant message, so truncating middle turns is safe.
    """
    if len(history) <= max_pairs * 2:
        return history
    first_pair = history[:2]                        # original notes + V1
    recent = history[-(max_pairs - 1) * 2:]         # last N-1 revision pairs
    return first_pair + recent


def search_semantic_scholar(query: str, max_results: int = 5) -> list[dict]:
    """Fetch related papers from Semantic Scholar (free, no auth required)."""
    try:
        resp = requests.get(
            "https://api.semanticscholar.org/graph/v1/paper/search",
            params={"query": query, "limit": max_results, "fields": "title,abstract,year,authors"},
            timeout=12,
        )
        resp.raise_for_status()
        return resp.json().get("data", [])
    except Exception:
        return []


_COUNCIL_BASE = """\
You are an AI Academic Council — four expert reviewers examining a research paper draft.
Speak as each reviewer in turn, then give a consensus.

**[Methods & Logic Reviewer]**
Assess logical flow, argument structure, and methodological claims.
Flag unsupported leaps, circular reasoning, or missing caveats. Be specific — quote the draft.

**[Literature & Evidence Reviewer]**
Cross-reference the draft's claims against the related papers provided.
Note: what aligns with prior work, what contradicts it, what needs a citation but has none.
If no related papers were found, flag claims that are likely to need literature support.

**[Academic Voice Reviewer]**
Flag AI-sounding phrases, over-hedging, passive voice overuse, and structural problems.
Rewrite the worst 2–3 sentences concretely.

**[Agricultural Economics Domain Reviewer]**
Evaluate the draft through the lens of applied agricultural economics.
Flag: vague cost/return language that lacks $/cwt or $/head units; break-even claims without \
explicit formula exposition; enterprise budget line items that are missing or implausible; \
feedlot management assumptions that contradict standard industry practice; policy implications \
that overreach the dataset; and missing citations to AEPP, JARE, Agribusiness, or AJAE literature. \
Quote the draft directly. Suggest the specific numbers or citations that are missing.

**[Council Consensus]**
Exactly 3 highest-priority actionable changes, numbered. No generic praise.

Core rules: quote actual phrases from the draft. Do not restate the draft. Be direct.\
"""


def build_council_prompt(council_rules: list | None, council_memory: str) -> str:
    """Inject standing instructions and accumulated memory into the council prompt."""
    prompt = _COUNCIL_BASE
    if council_rules:
        prompt += "\n\nSTANDING INSTRUCTIONS FROM AUTHOR — follow these in every review:\n"
        for i, r in enumerate(council_rules, 1):
            prompt += f"{i}. {r['text']}\n"
    if council_memory and council_memory.strip():
        prompt += (
            "\n\nCOUNCIL MEMORY — patterns observed across previous reviews of this author's work:\n"
            + council_memory.strip()
            + "\nApply these patterns to sharpen the current review — call out recurring issues immediately."
        )
    return prompt


def update_council_memory(review_text: str, session: dict,
                          client: anthropic.Anthropic, current_memory: str) -> str:
    """Ask Haiku to extract new patterns from a review and append them to memory."""
    paper_title = session.get("paper_title", "")
    section = session.get("section", "")
    prompt = (
        f"You maintain the memory of an AI academic review council.\n\n"
        f"EXISTING MEMORY:\n{current_memory.strip() or '(none yet)'}\n\n"
        f"A new review was just completed for the {section} section of \"{paper_title}\":\n\n"
        f"{review_text[:2500]}\n\n"
        f"Extract 1–3 specific, recurring patterns about this author's writing tendencies or "
        f"this paper's weaknesses that should inform future reviews. "
        f"If nothing new was observed beyond what's already in memory, reply with exactly: NONE"
        f"\n\nBe concise bullet points only. No preamble."
    )
    resp = client.messages.create(
        model=HAIKU_MODEL,
        max_tokens=250,
        messages=[{"role": "user", "content": prompt}],
    )
    new_patterns = resp.content[0].text.strip()
    if not new_patterns or new_patterns.upper() == "NONE":
        return current_memory
    timestamp = datetime.now().strftime("%Y-%m-%d")
    entry = f"\n\n[{timestamp} — {paper_title}, {section}]\n{new_patterns}"
    return (current_memory or "") + entry


def run_council_review(draft: str, session: dict, client: anthropic.Anthropic,
                       research_notes: str = "",
                       council_rules: list | None = None,
                       council_memory: str = "") -> str:
    """Run the AI Council review: Semantic Scholar search + three-voice critique."""
    paper_title = session.get("paper_title", "")
    section = session.get("section", "")
    ctx = session.get("paper_context", {})

    query_parts = [paper_title, section]
    if ctx.get("research_question"):
        query_parts.append(ctx["research_question"])
    if ctx.get("models_methods"):
        query_parts.append(ctx["models_methods"])
    query = " ".join(p for p in query_parts if p.strip())
    # Anchor to domain so Semantic Scholar returns relevant ag-econ papers
    if not any(kw in query.lower() for kw in ("feedlot", "beef cattle", "cattle", "livestock")):
        query = "feedlot economics beef cattle " + query

    related = search_semantic_scholar(query, max_results=5)

    if related:
        papers_block = "\nRELATED PAPERS FOUND (Semantic Scholar):\n"
        for i, p in enumerate(related, 1):
            authors = ", ".join(a.get("name", "") for a in p.get("authors", [])[:3])
            year = p.get("year", "n.d.")
            title = p.get("title", "Untitled")
            abstract = (p.get("abstract") or "")[:350].strip()
            papers_block += f"\n[{i}] {title} ({authors}, {year})\n"
            if abstract:
                papers_block += f"    {abstract}…\n"
    else:
        papers_block = "\n[Semantic Scholar returned no results — review without literature cross-reference]\n"

    user_msg = (
        f'Review the **{section}** section of the paper titled "{paper_title}".\n'
        f"{papers_block}\n"
        f"--- DRAFT ---\n{draft}\n--- END DRAFT ---"
    )
    if research_notes and research_notes.strip():
        user_msg += f"\n\nADDITIONAL PAPER CONTEXT (from author):\n{research_notes.strip()[:2000]}"

    system = build_council_prompt(council_rules, council_memory)
    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=system,
        messages=[{"role": "user", "content": user_msg}],
    )
    return response.content[0].text


def _rule_is_duplicate(new_rule: str, existing_rules: list) -> bool:
    """Return True if new_rule overlaps too heavily with any existing rule.

    Uses word-level Jaccard similarity: if the overlap exceeds 60 % of the
    smaller rule's vocabulary the rule is considered a duplicate.
    """
    def _words(text: str) -> set:
        return set(re.sub(r"[^a-z0-9 ]", " ", text.lower()).split())

    new_words = _words(new_rule)
    if not new_words:
        return False
    for r in existing_rules:
        existing_words = _words(r.get("text", ""))
        if not existing_words:
            continue
        intersection = len(new_words & existing_words)
        smaller = min(len(new_words), len(existing_words))
        if smaller and intersection / smaller >= 0.60:
            return True
    return False


def extract_rule_from_feedback(feedback: str, client: anthropic.Anthropic) -> str | None:
    """Use Haiku to extract a generalizable writing preference from revision feedback.

    Returns a rule string if feedback contains a reusable preference, else None.
    """
    prompt = (
        "A user gave feedback on a research paper draft. "
        "If the feedback reveals a generalizable writing preference, style rule, or pattern "
        "the author consistently cares about, extract it as one concise rule (1-2 sentences, "
        "starting with an imperative verb like 'Never', 'Always', 'Avoid', 'Keep').\n\n"
        f"FEEDBACK: {feedback}\n\n"
        "If the feedback is purely content-specific (add an example here, fix this number, "
        "change this claim) with no generalizable lesson, reply with exactly: NONE\n"
        "If there is a generalizable rule, reply with ONLY the rule. No explanation."
    )
    resp = client.messages.create(
        model=HAIKU_MODEL,
        max_tokens=80,
        messages=[{"role": "user", "content": prompt}],
    )
    result = resp.content[0].text.strip()
    if not result or result.upper().startswith("NONE"):
        return None
    return result


def parse_consensus_items(review_text: str) -> list[str]:
    """Extract the numbered items from [Council Consensus] in a review.

    Handles bold-formatted numbers like **1. Headline** explanation text,
    plain numbered items, and variations the council may produce.
    """
    idx = review_text.lower().find("[council consensus]")
    if idx == -1:
        return []
    section = review_text[idx:]

    # Strip all markdown bold/italic before parsing so numbers are bare
    section = re.sub(r'\*\*([^*\n]+)\*\*', r'\1', section)
    section = re.sub(r'\*([^*\n]+)\*', r'\1', section)

    # Match: optional whitespace, digit(s), dot or paren, space, then content
    # Stops at the next numbered item or end of string
    items = re.findall(
        r'(?:^|\n)[ \t]*\d+[.)]\s+(.+?)(?=\n[ \t]*\d+[.)]|\Z)',
        section,
        re.DOTALL,
    )

    cleaned = []
    for text in items:
        # Collapse whitespace
        t = ' '.join(text.split())
        if not t:
            continue
        # Use just the first sentence as the checkbox label (keep it short)
        first = re.split(r'\.\s+[A-Z]', t)[0].strip()
        if first and not first.endswith('.'):
            first += '.'
        cleaned.append(first[:250] if first else t[:250])
    return cleaned


def find_related_sessions(current_id: str, current_title: str, all_sessions: dict) -> list[dict]:
    """Return other sessions sharing the same paper title that have a current draft."""
    title_lower = current_title.lower().strip()
    section_order = {s: i for i, s in enumerate(SECTIONS)}
    related = [
        s for sid, s in all_sessions.items()
        if sid != current_id
        and s.get("paper_title", "").lower().strip() == title_lower
        and s.get("current_draft")
    ]
    related.sort(key=lambda s: section_order.get(s.get("section", ""), 99))
    return related


def stream_claude(conversation_history: list, session: dict, client: anthropic.Anthropic,
                  rules: list | None = None, research_notes: str = "",
                  related_sessions: list | None = None):
    """Stream Claude's response token by token. Raises on failure."""
    with client.messages.stream(
        model=MODEL,
        max_tokens=8192,
        system=build_system_prompt(session, rules=rules, research_notes=research_notes,
                                   related_sessions=related_sessions),
        messages=truncate_history(conversation_history),
    ) as stream:
        for text in stream.text_stream:
            yield text


# ── UI: Sidebar ────────────────────────────────────────────────────────────────

def render_sidebar(sessions: dict) -> None:
    with st.sidebar:
        st.title("Paper Agent")
        if st.button("＋ New Session", use_container_width=True, type="primary"):
            st.session_state.current_session_id = None
            st.rerun()

        # ── Learned preferences (auto-extracted from your feedback) ──
        st.divider()
        rules = st.session_state.get("rules", [])
        lp_label = f"Learned Preferences ({len(rules)})" if rules else "Learned Preferences"
        with st.expander(lp_label, expanded=False):
            if rules:
                st.caption(
                    "Extracted automatically from your revision feedback. "
                    "Applied to every draft. Remove any that don't fit."
                )
                for i, rule in enumerate(rules):
                    col_text, col_del = st.columns([5, 1])
                    with col_text:
                        src = " _(auto)_" if rule.get("auto") else ""
                        st.markdown(f"**{i+1}.** {rule['text']}{src}")
                        st.caption(f"Learned {rule['added_at'][:10]}")
                    with col_del:
                        if st.button("✕", key=f"del_rule_{i}", help="Remove"):
                            rules.pop(i)
                            save_rules(rules)
                            st.session_state.rules = rules
                            st.rerun()
            else:
                st.caption(
                    "Nothing learned yet. Give the agent revision feedback and "
                    "it will extract your preferences automatically."
                )

        # ── Persistent research notes ──
        st.divider()
        research_notes = st.session_state.get("research_notes", "")
        rn_label = "Research Notes (saved)" if research_notes.strip() else "Research Notes — paste your research here"
        with st.expander(rn_label, expanded=(not research_notes.strip())):
            st.caption(
                "Paste anything you never want to re-explain: paper details, formulas, "
                "dataset stats, key findings. Saved globally — available in every session, forever."
            )
            edited = st.text_area(
                "research_notes_input",
                value=research_notes,
                height=260,
                label_visibility="collapsed",
                placeholder=(
                    "e.g.:\n"
                    "Paper: FEDVT — Excel/VBA feedlot cost tool\n"
                    "TC = C_feed + C_feeder + C_other + C_interest + C_fixed + C_labor\n"
                    "BEP = (TC / TFW) × 100 + FB\n"
                    "Default herd: 2000 head, 600 lb feeder @ $250/cwt, finish 1400 lb...\n"
                ),
                key="research_notes_textarea",
            )
            if st.button("Save Research Notes", type="primary", use_container_width=True, key="save_rn"):
                save_research_notes(edited)
                st.session_state.research_notes = edited
                st.success("Saved — injected into every prompt from now on.")

        # ── Council settings ──
        st.divider()
        council_rules = st.session_state.get("council_rules", [])
        council_memory = st.session_state.get("council_memory", "")
        auto_review = st.session_state.get("auto_review", False)
        n_cr = len(council_rules)
        memory_entries = council_memory.count("[20") if council_memory else 0
        cr_label = f"Council Settings ({n_cr} instructions · {memory_entries} memories)" if (n_cr or memory_entries) else "Council Settings"
        with st.expander(cr_label, expanded=False):
            new_auto = st.toggle("Auto-review every draft", value=auto_review,
                                 help="Automatically run council review after each generation or revision")
            if new_auto != auto_review:
                st.session_state.auto_review = new_auto

            st.caption("Standing instructions — tell the council what to focus on:")
            if council_rules:
                for i, rule in enumerate(council_rules):
                    col_t, col_x = st.columns([5, 1])
                    with col_t:
                        st.markdown(f"**{i+1}.** {rule['text']}")
                        st.caption(f"Added {rule['added_at'][:10]}")
                    with col_x:
                        if st.button("✕", key=f"del_cr_{i}", help="Remove"):
                            council_rules.pop(i)
                            save_council_rules(council_rules)
                            st.session_state.council_rules = council_rules
                            st.rerun()
                st.divider()
            with st.form("add_council_rule_form"):
                new_cr = st.text_area(
                    "new_council_rule",
                    placeholder=(
                        'e.g., "Focus on missing citations in feedlot economics literature"\n'
                        'e.g., "Be harsh on passive voice — this is a top-tier journal target"\n'
                        'e.g., "Always check if the Methods section quantifies uncertainty"'
                    ),
                    height=90,
                    label_visibility="collapsed",
                )
                if st.form_submit_button("Add Instruction", type="primary", use_container_width=True):
                    if new_cr.strip():
                        council_rules.append({"text": new_cr.strip(), "added_at": datetime.now().isoformat()})
                        save_council_rules(council_rules)
                        st.session_state.council_rules = council_rules
                        st.rerun()

            if council_memory and council_memory.strip():
                st.divider()
                st.caption(f"Council memory ({memory_entries} session{'s' if memory_entries != 1 else ''}):")
                with st.expander("View memory", expanded=False):
                    st.text(council_memory)
                if st.button("Clear memory", use_container_width=True):
                    save_council_memory("")
                    st.session_state.council_memory = ""
                    st.rerun()
            else:
                st.caption("_Memory builds automatically after each review._")

        st.divider()
        st.subheader("Sessions")

        if not sessions:
            st.caption("No sessions yet.")
            return

        for sid, s in sorted(sessions.items(), key=lambda x: x[1].get("last_modified", x[1]["timestamp"]), reverse=True):
            title = s["paper_title"]
            title_short = (title[:22] + "…") if len(title) > 25 else title
            date = s.get("last_modified", s["timestamp"])[:10]
            n_versions = len(s.get("drafts", []))
            n_docs = len(s.get("documents", []))
            meta_parts = [s["section"]]
            if n_versions:
                meta_parts.append(f"{n_versions}v")
            if n_docs:
                meta_parts.append(f"{n_docs} doc{'s' if n_docs > 1 else ''}")
            meta = " · ".join(meta_parts)
            label = f"{title_short}\n{meta} · {date}"
            is_active = st.session_state.get("current_session_id") == sid

            col_btn, col_del = st.columns([5, 1])
            with col_btn:
                if st.button(label, key=f"btn_{sid}", use_container_width=True,
                             type="primary" if is_active else "secondary"):
                    st.session_state.current_session_id = sid
                    st.rerun()
            with col_del:
                if st.button("✕", key=f"del_{sid}", help="Delete this session"):
                    del st.session_state.sessions[sid]
                    if st.session_state.get("current_session_id") == sid:
                        st.session_state.current_session_id = None
                    save_sessions(st.session_state.sessions)
                    st.rerun()


# ── UI: New session form ───────────────────────────────────────────────────────

def render_new_session_form() -> None:
    st.title("Research Paper Writing Agent")
    st.markdown(
        "A style-aware writing partner that drafts and iterates on research paper sections "
        "in your voice. Sessions are saved automatically — pick up exactly where you left off."
    )
    st.divider()

    with st.form("new_session_form"):
        paper_title = st.text_input(
            "Paper Title",
            placeholder="e.g., Scaling Laws for Neural Language Models",
        )
        section = st.selectbox("Section to Draft", SECTIONS)
        submitted = st.form_submit_button("Start Session →", type="primary")

        if submitted:
            if not paper_title.strip():
                st.error("Enter a paper title to continue.")
            else:
                sid, session = new_session(paper_title.strip(), section)
                st.session_state.sessions[sid] = session
                save_sessions(st.session_state.sessions)
                st.session_state.current_session_id = sid
                st.rerun()


# ── UI: Paper context panel ────────────────────────────────────────────────────

def render_paper_context_panel(session: dict) -> None:
    """Editable paper context stored per session."""
    ctx = session.setdefault("paper_context", {key: "" for key, *_ in _CONTEXT_FIELDS})
    filled = sum(1 for v in ctx.values() if v.strip())
    label = (
        f"Paper Context ({filled}/{len(_CONTEXT_FIELDS)} fields filled)"
        if filled else
        "Paper Context — tell the agent about your paper"
    )

    with st.expander(label, expanded=(filled == 0)):
        with st.form(f"ctx_form_{session['session_id']}"):
            new_vals = {}
            for key, label_text, placeholder in _CONTEXT_FIELDS:
                new_vals[key] = st.text_input(
                    label_text,
                    value=ctx.get(key, ""),
                    placeholder=placeholder,
                )
            if st.form_submit_button("Save Context", type="primary"):
                for key, val in new_vals.items():
                    ctx[key] = val
                touch_session(session)
                save_sessions(st.session_state.sessions)
                st.success("Context saved — used on the next generation.")


# ── UI: Document panel ─────────────────────────────────────────────────────────

def render_document_panel(session: dict) -> None:
    """Render the document attachment section."""
    docs = session.setdefault("documents", [])
    n = len(docs)
    label = f"Reference Documents ({n} attached)" if n else "Reference Documents"

    with st.expander(label, expanded=(n > 0)):
        # Show attached docs
        if docs:
            for i, doc in enumerate(docs):
                col_name, col_remove = st.columns([5, 1])
                with col_name:
                    chars = len(doc["content"])
                    st.markdown(f"**{doc['name']}** — {chars:,} chars · added {doc['added_at'][:10]}")
                with col_remove:
                    if st.button("✕", key=f"remove_doc_{i}", help="Remove this document"):
                        docs.pop(i)
                        touch_session(session)
                        save_sessions(st.session_state.sessions)
                        st.rerun()
            st.divider()

        # Add by URL
        st.markdown("**Add from URL** (Google Docs, Google Drive, any web page or PDF link)")
        col_url, col_add_url = st.columns([4, 1])
        with col_url:
            url_input = st.text_input(
                "url_input",
                placeholder="https://docs.google.com/document/d/... or any URL",
                label_visibility="collapsed",
                key="doc_url_input",
            )
        with col_add_url:
            if st.button("Add →", key="add_url_btn", use_container_width=True):
                if not url_input.strip():
                    st.error("Paste a URL first.")
                else:
                    with st.spinner("Fetching…"):
                        try:
                            name, content = fetch_url_content(url_input.strip())
                            if not content.strip():
                                st.error("Could not extract text from that URL.")
                            else:
                                docs.append({
                                    "name": name,
                                    "source": url_input.strip(),
                                    "content": content[:DOC_CONTENT_LIMIT],
                                    "added_at": datetime.now().isoformat(),
                                })
                                touch_session(session)
                                save_sessions(st.session_state.sessions)
                                st.success(f"Added: {name} ({len(content):,} chars)")
                                st.rerun()
                        except Exception as e:
                            st.error(f"Failed to fetch URL: {e}")

        # Add by file upload (all supported types)
        st.markdown("**Upload a file** (PDF, DOCX, TXT, MD, CSV, XLSX)")
        uploaded_file = st.file_uploader(
            "upload_doc",
            type=["pdf", "docx", "txt", "md", "tex", "rst", "csv", "xlsx", "xls"],
            label_visibility="collapsed",
            key="doc_file_uploader",
        )
        if uploaded_file is not None:
            already_added = any(d["name"] == uploaded_file.name for d in docs)
            if already_added:
                st.caption(f"✓ {uploaded_file.name} is already attached.")
            else:
                if st.button(f"Add \"{uploaded_file.name}\"", key="confirm_upload_btn"):
                    with st.spinner("Extracting text…"):
                        try:
                            content = extract_file_text(uploaded_file.read(), uploaded_file.name)
                            if not content.strip():
                                st.error("Could not extract text from this file.")
                            else:
                                docs.append({
                                    "name": uploaded_file.name,
                                    "source": "uploaded file",
                                    "content": content[:DOC_CONTENT_LIMIT],
                                    "added_at": datetime.now().isoformat(),
                                })
                                touch_session(session)
                                save_sessions(st.session_state.sessions)
                                st.success(f"Added: {uploaded_file.name} ({len(content):,} chars)")
                                st.rerun()
                        except Exception as e:
                            st.error(f"Failed to read file: {e}")

        # Add by local file path (fast for large files — no browser upload needed)
        st.markdown("**Add from local path** (paste the full file path on your machine)")
        col_path, col_add_path = st.columns([4, 1])
        with col_path:
            path_input = st.text_input(
                "local_path_input",
                placeholder=r"e.g.  C:\Users\shiva\Desktop\data.xlsx  or  /home/user/notes.txt",
                label_visibility="collapsed",
                key="doc_path_input",
            )
        with col_add_path:
            if st.button("Add →", key="add_path_btn", use_container_width=True):
                if not path_input.strip():
                    st.error("Enter a file path first.")
                else:
                    with st.spinner("Reading…"):
                        try:
                            name, content = read_local_path(path_input)
                            if not content.strip():
                                st.error("File is empty or no text could be extracted.")
                            else:
                                already = any(d["name"] == name for d in docs)
                                if already:
                                    st.warning(f"{name} is already attached.")
                                else:
                                    docs.append({
                                        "name": name,
                                        "source": path_input.strip(),
                                        "content": content[:DOC_CONTENT_LIMIT],
                                        "added_at": datetime.now().isoformat(),
                                    })
                                    touch_session(session)
                                    save_sessions(st.session_state.sessions)
                                    st.success(f"Added: {name} ({len(content):,} chars)")
                                    st.rerun()
                        except Exception as e:
                            st.error(f"Could not read file: {e}")

        if docs:
            st.caption(
                "Documents are saved to this session and automatically included as context "
                "on every generation and revision."
            )


# ── UI: Download buttons ───────────────────────────────────────────────────────

def render_download_buttons(draft_content: str, paper_title: str, section: str,
                            version: str, key_prefix: str) -> None:
    """Render TXT, DOCX, and PDF download buttons stacked in a column."""
    safe_title = re.sub(r"[^\w\s-]", "", paper_title)[:40].strip()
    base = f"{safe_title}_{section}_{version}"

    st.download_button(
        "⬇ TXT",
        data=draft_content,
        file_name=f"{base}.txt",
        mime="text/plain",
        use_container_width=True,
        key=f"{key_prefix}_txt",
    )

    if DOCX_SUPPORT:
        try:
            docx_bytes = generate_docx(draft_content, paper_title, section)
            st.download_button(
                "⬇ DOCX",
                data=docx_bytes,
                file_name=f"{base}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key=f"{key_prefix}_docx",
                help="Best for Google Docs — import via File → Open",
            )
        except Exception:
            pass
    else:
        st.caption("_Install python-docx for DOCX_")

    if FPDF_SUPPORT:
        try:
            pdf_bytes = generate_pdf(draft_content, paper_title, section)
            st.download_button(
                "⬇ PDF",
                data=pdf_bytes,
                file_name=f"{base}.pdf",
                mime="application/pdf",
                use_container_width=True,
                key=f"{key_prefix}_pdf",
            )
        except Exception:
            pass
    else:
        st.caption("_Install fpdf2 for PDF_")


# ── UI: Active session ─────────────────────────────────────────────────────────

def render_active_session(session: dict, session_id: str, client: anthropic.Anthropic) -> None:
    # Backfill paper_context for sessions created before this field existed
    session.setdefault("paper_context", {key: "" for key, *_ in _CONTEXT_FIELDS})

    # Cross-session context: other sections of the same paper
    related_sessions = find_related_sessions(
        session_id, session.get("paper_title", ""), st.session_state.sessions
    )

    has_draft = bool(session["current_draft"])
    current_version = session["drafts"][-1]["version"] if session["drafts"] else ""
    n_turns = len(session["conversation_history"]) // 2
    last_saved = session.get("last_modified", session["timestamp"])[:16].replace("T", " ")

    # ── Header ──
    st.markdown(f"## {session['paper_title']}")
    col1, col2, col3, col4 = st.columns(4)
    with col1: st.markdown(f"**Section:** {session['section']}")
    with col2: st.markdown(f"**Session:** `{session_id}`")
    with col3: st.markdown(f"**Turns:** {n_turns}")
    with col4: st.markdown(f"**Saved:** {last_saved}")
    st.divider()

    # Backfill timestamp on old draft entries that predate this field
    for draft in session["drafts"]:
        draft.setdefault("timestamp", session["timestamp"])

    # ── Tabs ──
    write_tab, history_tab, setup_tab = st.tabs(["Write", f"History ({len(session['drafts'])} drafts)", "Setup"])

    # ════════════════════════════════════════════════
    # WRITE TAB
    # ════════════════════════════════════════════════
    with write_tab:
        left, right = st.columns([1, 1], gap="large")

        with left:
            if not has_draft:
                st.subheader("Your Notes")
                notes = st.text_area(
                    "notes_input",
                    height=320,
                    placeholder=(
                        "Paste raw notes, bullet points, or a rough outline:\n\n"
                        "• Main claim: model X outperforms baselines by 4.2% BLEU\n"
                        "• Why: larger receptive field, no recurrence bottleneck\n"
                        "• Caveat: tested on English-German only\n"
                        "• Include comparison table reference"
                    ),
                    label_visibility="collapsed",
                )
                if st.button("Generate Draft →", type="primary", use_container_width=True):
                    if not notes.strip():
                        st.error("Paste some notes first.")
                    else:
                        n_docs = len(session.get("documents", []))
                        doc_note = f" (using {n_docs} attached document{'s' if n_docs > 1 else ''})" if n_docs else ""
                        user_msg = (
                            f"Draft the {session['section']} section for my paper "
                            f"titled '{session['paper_title']}'.\n\nNotes:\n\n{notes}"
                        )
                        session["conversation_history"].append({"role": "user", "content": user_msg})
                        with right:
                            st.subheader(f"Generating Draft{doc_note}…")
                            placeholder = st.empty()
                        collected = []
                        try:
                            for chunk in stream_claude(session["conversation_history"], session, client, rules=st.session_state.get("rules"), research_notes=st.session_state.get("research_notes", ""), related_sessions=related_sessions):
                                collected.append(chunk)
                                placeholder.markdown("".join(collected))
                        except Exception as e:
                            session["conversation_history"].pop()
                            st.error(f"Generation failed: {e}")
                            st.stop()
                        draft = "".join(collected)
                        session["conversation_history"].append({"role": "assistant", "content": draft})
                        session["current_draft"] = draft
                        session["drafts"].append({"version": "V1", "content": draft, "timestamp": datetime.now().isoformat()})
                        touch_session(session)
                        save_sessions(st.session_state.sessions)
                        st.rerun()

            else:
                st.subheader(f"Revise {current_version}")
                feedback = st.text_area(
                    "feedback_input",
                    height=280,
                    placeholder=(
                        "Describe what to change:\n\n"
                        '"The second paragraph hedges too much—make it direct."\n'
                        '"Cut the last sentence, it repeats the point above."\n'
                        '"Add a concrete example to the third claim."\n'
                        '"The methods description is too vague—be more specific."'
                    ),
                    label_visibility="collapsed",
                )
                col_revise, col_reset = st.columns([2, 1])
                with col_revise:
                    if st.button("Revise Draft →", type="primary", use_container_width=True):
                        if not feedback.strip():
                            st.error("Enter feedback first.")
                        else:
                            user_msg = (
                                f"Revise the {session['section']} section draft based on this feedback:\n\n"
                                f"{feedback}\n\n"
                                f"Preserve all parts not addressed in the feedback exactly as they are."
                            )
                            session["conversation_history"].append({"role": "user", "content": user_msg})
                            with right:
                                st.subheader(f"Revising {current_version}…")
                                placeholder = st.empty()
                            collected = []
                            try:
                                for chunk in stream_claude(session["conversation_history"], session, client, rules=st.session_state.get("rules"), research_notes=st.session_state.get("research_notes", ""), related_sessions=related_sessions):
                                    collected.append(chunk)
                                    placeholder.markdown("".join(collected))
                            except Exception as e:
                                session["conversation_history"].pop()
                                st.error(f"Revision failed: {e}")
                                st.stop()
                            revised = "".join(collected)
                            session["conversation_history"].append({"role": "assistant", "content": revised})
                            session["current_draft"] = revised
                            session["iteration_notes"].append(feedback)
                            new_ver = f"V{len(session['drafts']) + 1}"
                            session["drafts"].append({"version": new_ver, "content": revised, "timestamp": datetime.now().isoformat()})
                            touch_session(session)
                            save_sessions(st.session_state.sessions)
                            # Auto-learn from feedback: extract generalizable preference
                            try:
                                learned = extract_rule_from_feedback(feedback, client)
                                if learned:
                                    rules = st.session_state.get("rules", [])
                                    if not _rule_is_duplicate(learned, rules):
                                        rules.append({
                                            "text": learned,
                                            "added_at": datetime.now().isoformat(),
                                            "auto": True,
                                        })
                                        save_rules(rules)
                                        st.session_state.rules = rules
                            except Exception:
                                pass  # Don't block the UI if rule extraction fails
                            st.rerun()

                with col_reset:
                    confirm_key = f"confirm_reset_{session_id}"
                    if st.session_state.get(confirm_key):
                        if st.button("⚠ Confirm Reset", key="confirm_yes", use_container_width=True, type="primary"):
                            session["current_draft"] = ""
                            session["conversation_history"] = []
                            session["drafts"] = []
                            session["iteration_notes"] = []
                            session["council_reviews"] = {}
                            st.session_state[confirm_key] = False
                            touch_session(session)
                            save_sessions(st.session_state.sessions)
                            st.rerun()
                        if st.button("Cancel", key="confirm_no", use_container_width=True):
                            st.session_state[confirm_key] = False
                            st.rerun()
                    else:
                        if st.button("Start Over", use_container_width=True):
                            st.session_state[confirm_key] = True
                            st.rerun()

        with right:
            if has_draft:
                word_count = len(session["current_draft"].split())
                col_hdr, col_dl = st.columns([3, 1])
                with col_hdr:
                    st.subheader(f"Current Draft — {current_version}")
                    st.caption(f"{word_count:,} words")
                with col_dl:
                    render_download_buttons(
                        session["current_draft"],
                        session["paper_title"],
                        session["section"],
                        current_version,
                        f"cur_{session_id}",
                    )
                with st.container(border=True):
                    st.markdown(session["current_draft"])

                # ── Council Review ──
                st.divider()
                reviews = session.setdefault("council_reviews", {})
                existing = reviews.get(current_version)

                def _do_council_review():
                    """Run review, update memory, persist — shared by auto and manual triggers."""
                    rev = run_council_review(
                        session["current_draft"], session, client,
                        st.session_state.get("research_notes", ""),
                        council_rules=st.session_state.get("council_rules"),
                        council_memory=st.session_state.get("council_memory", ""),
                    )
                    reviews[current_version] = rev
                    # Update council memory with patterns from this review
                    new_mem = update_council_memory(
                        rev, session, client,
                        st.session_state.get("council_memory", ""),
                    )
                    if new_mem != st.session_state.get("council_memory", ""):
                        save_council_memory(new_mem)
                        st.session_state.council_memory = new_mem
                    touch_session(session)
                    save_sessions(st.session_state.sessions)

                # Auto-review: fires once per version when toggle is on
                if not existing and st.session_state.get("auto_review", False):
                    with st.spinner("Council auto-reviewing draft…"):
                        _do_council_review()
                        st.rerun()

                if existing:
                    col_cr_hdr, col_cr_rerun = st.columns([4, 1])
                    with col_cr_hdr:
                        st.subheader("Council Review")
                    with col_cr_rerun:
                        if st.button("↺ Re-run", key="rerun_review", use_container_width=True):
                            with st.spinner("Re-reviewing…"):
                                _do_council_review()
                                st.rerun()
                    with st.expander("Full review", expanded=False):
                        st.markdown(existing)

                    # ── Accept-critiques workflow ──
                    consensus_items = parse_consensus_items(existing)
                    if consensus_items:
                        st.caption("Select council critiques to accept, then remake the draft:")
                        accepted = []
                        for i, item in enumerate(consensus_items):
                            if st.checkbox(item, key=f"cr_{session_id}_{current_version}_{i}"):
                                accepted.append(item)
                        if accepted:
                            if st.button("Accept & Remake Draft", key="accept_remake",
                                         type="primary", use_container_width=True):
                                # Add accepted critiques to learned preferences
                                for critique in accepted:
                                    try:
                                        rule = extract_rule_from_feedback(critique, client)
                                        if rule:
                                            rl = st.session_state.get("rules", [])
                                            if not _rule_is_duplicate(rule, rl):
                                                rl.append({"text": rule, "added_at": datetime.now().isoformat(), "auto": True})
                                                save_rules(rl)
                                                st.session_state.rules = rl
                                    except Exception:
                                        pass
                                apply_msg = (
                                    f"Revise the {session['section']} section to address these accepted critiques:\n\n"
                                    + "\n".join(f"{j+1}. {c}" for j, c in enumerate(accepted))
                                    + "\n\nPreserve all parts not related to these critiques exactly as they are."
                                )
                                session["conversation_history"].append({"role": "user", "content": apply_msg})
                                with st.spinner("Remaking draft with accepted critiques…"):
                                    placeholder = st.empty()
                                    collected = []
                                    try:
                                        for chunk in stream_claude(
                                            session["conversation_history"], session, client,
                                            rules=st.session_state.get("rules"),
                                            research_notes=st.session_state.get("research_notes", ""),
                                            related_sessions=related_sessions,
                                        ):
                                            collected.append(chunk)
                                            placeholder.markdown("".join(collected))
                                    except Exception as e:
                                        session["conversation_history"].pop()
                                        st.error(f"Remake failed: {e}")
                                        st.stop()
                                    revised = "".join(collected)
                                    session["conversation_history"].append({"role": "assistant", "content": revised})
                                    session["current_draft"] = revised
                                    session["iteration_notes"].append(f"Council: {'; '.join(accepted[:2])}")
                                    new_ver = f"V{len(session['drafts']) + 1}"
                                    session["drafts"].append({"version": new_ver, "content": revised, "timestamp": datetime.now().isoformat()})
                                    touch_session(session)
                                    save_sessions(st.session_state.sessions)
                                    st.rerun()
                else:
                    if st.button(
                        "Run Council Review",
                        key="run_review",
                        use_container_width=True,
                        type="primary",
                        help="Three AI reviewers cross-reference your draft against Semantic Scholar research",
                    ):
                        with st.spinner("Searching Semantic Scholar and reviewing draft…"):
                            _do_council_review()
                            st.rerun()
            else:
                st.subheader("Draft")
                st.info("Your draft will appear here after generation.")

    # ════════════════════════════════════════════════
    # SETUP TAB
    # ════════════════════════════════════════════════
    with setup_tab:
        render_paper_context_panel(session)
        st.divider()
        render_document_panel(session)

    # ════════════════════════════════════════════════
    # HISTORY TAB
    # ════════════════════════════════════════════════
    with history_tab:
        if not session["drafts"]:
            st.info("No drafts yet. Generate one in the Write tab.")
        else:
            safe_title = re.sub(r"[^\w\s-]", "", session["paper_title"])[:40].strip()
            for idx, draft in reversed(list(enumerate(session["drafts"]))):
                is_current = (idx == len(session["drafts"]) - 1)
                version_label = draft["version"] + (" — current" if is_current else "")
                word_count = len(draft["content"].split())
                drafted_at = draft.get("timestamp", session["timestamp"])[:16].replace("T", " ")

                col_meta, col_dl = st.columns([3, 1])
                with col_meta:
                    st.subheader(version_label)
                    st.caption(f"{word_count:,} words · {drafted_at}")
                    if idx == 0:
                        st.caption("_Initial generation from notes_")
                    elif idx - 1 < len(session["iteration_notes"]):
                        st.caption(f"Feedback: _{session['iteration_notes'][idx - 1]}_")
                with col_dl:
                    render_download_buttons(
                        draft["content"],
                        session["paper_title"],
                        session["section"],
                        draft["version"],
                        f"hist_{idx}",
                    )
                    if not is_current:
                        if st.button("↩ Restore", key=f"restore_{idx}",
                                     use_container_width=True,
                                     help="Set this version as the current draft"):
                            session["current_draft"] = draft["content"]
                            touch_session(session)
                            save_sessions(st.session_state.sessions)
                            st.rerun()
                with st.container(border=True):
                    st.markdown(draft["content"])
                stored_review = session.get("council_reviews", {}).get(draft["version"])
                if stored_review:
                    with st.expander("Council Review", expanded=False):
                        st.markdown(stored_review)
                if idx > 0:
                    st.divider()


# ── Entry point ────────────────────────────────────────────────────────────────

def main() -> None:
    st.set_page_config(
        page_title="Paper Writing Agent",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    try:
        client = anthropic.Anthropic()
    except Exception:
        st.error(
            "**ANTHROPIC_API_KEY not set.**\n\n"
            "In PowerShell:\n```\n$env:ANTHROPIC_API_KEY = 'sk-ant-...'\nstreamlit run app.py\n```"
        )
        st.stop()

    if "sessions" not in st.session_state:
        st.session_state.sessions = load_sessions()
    if "rules" not in st.session_state:
        st.session_state.rules = load_rules()
    if "council_rules" not in st.session_state:
        st.session_state.council_rules = load_council_rules()
    if "council_memory" not in st.session_state:
        st.session_state.council_memory = load_council_memory()
    if "research_notes" not in st.session_state:
        st.session_state.research_notes = load_research_notes()
    if "current_session_id" not in st.session_state:
        st.session_state.current_session_id = None
    if "auto_review" not in st.session_state:
        st.session_state.auto_review = False

    render_sidebar(st.session_state.sessions)

    current_id = st.session_state.current_session_id
    if current_id is None or current_id not in st.session_state.sessions:
        render_new_session_form()
    else:
        render_active_session(st.session_state.sessions[current_id], current_id, client)


if __name__ == "__main__":
    main()
